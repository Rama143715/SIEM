from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/AI_SIEM_Platform_Lab_Submission_Guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    if color:
      run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, header, bold=True, color="1F4D78")
        if widths:
            cell.width = Inches(widths[idx])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
            if widths:
                cells[idx].width = Inches(widths[idx])

    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    title_run = p.add_run(f"{title}: ")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(31, 77, 120)
    body_run = p.add_run(text)
    body_run.font.size = Pt(10)
    doc.add_paragraph()


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(89, 89, 89)

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("AI SIEM Platform Lab Submission Guide")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    add_footer(section)
    set_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("AI SIEM Platform")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Complete Lab Submission Guide, Project Information, Workflow, and Technology Stack")

    doc.add_paragraph()
    add_callout(
        doc,
        "Project summary",
        "A full-stack Security Operations Center prototype that collects security logs, detects threats with rules, creates alerts, supports incident tracking, and provides AI-assisted analysis with offline fallback mode.",
    )

    add_table(
        doc,
        ["Item", "Details"],
        [
            ["Application URL", "http://127.0.0.1:5173"],
            ["Database GUI", "http://127.0.0.1:8080"],
            ["Default user", "admin@siem.local / changeme123"],
            ["Submission purpose", "Demonstrate SIEM architecture, detection rules, alert handling, incidents, and database records."],
        ],
        widths=[1.7, 4.8],
    )

    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "The AI SIEM Platform is a lab-ready SIEM and SOC dashboard. It simulates how security teams collect logs, detect suspicious behavior, create alerts, and investigate incidents. The platform is designed for a real-world inspired demonstration while staying safe for a local lab machine."
    )
    add_bullets(
        doc,
        [
            "Central dashboard for threat monitoring and SOC metrics.",
            "Log ingestion through REST APIs, collector endpoint, and syslog UDP.",
            "Rule-based alerting for brute force, SQL injection, credential dumping, and suspicious login chains.",
            "Incident management for analyst investigation and response tracking.",
            "AI Analysis page with offline deterministic analysis when no external API key is configured.",
            "PostgreSQL database GUI for showing stored users, logs, alerts, rules, incidents, and audit records.",
        ],
    )

    doc.add_heading("2. Technology Stack", level=1)
    add_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "React 18, Vite, TailwindCSS, Recharts", "SOC dashboard, pages, charts, forms, and analyst workflow."],
            ["Backend", "Node.js, Express, Socket.IO", "Authentication, APIs, ingestion, alert processing, and real-time updates."],
            ["Database", "PostgreSQL 16", "Stores users, logs, alerts, incidents, rules, audit logs, and AI analysis history."],
            ["Cache/Queue", "Redis 7", "Supports rate limiting and background processing behavior."],
            ["Security", "JWT, bcrypt, Helmet, CORS, Joi", "Authentication, password hashing, API hardening, and validation."],
            ["Deployment", "Docker Compose", "Runs frontend, backend, Postgres, Redis, and Adminer locally."],
            ["Database GUI", "Adminer", "Browser-based GUI for viewing database records."],
            ["AI", "Anthropic SDK with offline fallback", "Online Claude integration when key exists; offline lab mode otherwise."],
        ],
        widths=[1.1, 2.1, 3.3],
    )

    doc.add_heading("3. System Architecture", level=1)
    add_table(
        doc,
        ["Component", "Port", "Role"],
        [
            ["Frontend container", "5173 -> 80", "Serves the React application through Nginx."],
            ["Backend container", "3001", "Provides REST API, auth, ingestion, rules, alerts, and Socket.IO."],
            ["PostgreSQL container", "5432", "Stores persistent application data."],
            ["Redis container", "6379", "Supports rate limits and service coordination."],
            ["Adminer container", "8080", "Database GUI for lab inspection."],
            ["Syslog UDP", "5514", "Receives syslog-style security events."],
        ],
        widths=[2.0, 1.2, 3.3],
    )
    add_callout(
        doc,
        "Architecture flow",
        "Browser -> React frontend -> Express backend API -> PostgreSQL/Redis. Security logs enter the backend, rules detect suspicious patterns, alerts are created, and analysts review results in the SOC interface.",
    )

    doc.add_heading("4. Application Workflow", level=1)
    add_steps(
        doc,
        [
            "User logs in with the admin account.",
            "Security events are ingested through API, collector, demo script, or syslog.",
            "Backend normalizes and stores events in PostgreSQL.",
            "Detection rules check each event or event pattern.",
            "Matching rules create alerts with severity and status.",
            "Dashboard, Logs, Alerts, Incidents, AI Analysis, Rules, and Settings pages allow analyst investigation.",
            "Resolved alerts and incident records show response completion.",
        ],
    )

    doc.add_heading("5. Main Application Modules", level=1)
    add_table(
        doc,
        ["Module", "What it shows", "How to explain it"],
        [
            ["Dashboard", "Threat metrics, open alerts, severity chart, source activity, asset health.", "This is the SOC overview screen."],
            ["Logs", "Ingested event records with search/filtering.", "Raw and normalized security evidence is stored here."],
            ["Alerts", "Rule detections with acknowledge and resolve actions.", "Open -> Acknowledge -> Resolve represents analyst workflow."],
            ["Incidents", "Investigation records and timelines.", "Analysts group related alerts/logs into a case."],
            ["AI Analysis", "Offline or API-backed triage notes.", "Summarizes attack type, indicators, MITRE hints, and response steps."],
            ["Rules", "Detection rule configuration.", "Rules convert suspicious logs into alerts."],
            ["Settings", "Access details, login/logout, password change, project about.", "User account and project explanation area."],
        ],
        widths=[1.3, 2.6, 2.6],
    )

    doc.add_heading("6. How to Start the Project", level=1)
    add_steps(
        doc,
        [
            'Open PowerShell and run: cd "C:\\Users\\bpava\\OneDrive\\Documents\\New project\\Wazuh-SIEM-Security-Lab\\siem-platform"',
            "Start Docker Desktop and wait until the engine is running.",
            "Run: docker compose up -d --build",
            "Check status: docker compose ps",
            "Open the application: http://127.0.0.1:5173",
            "Login with admin@siem.local and password changeme123.",
        ],
    )
    add_table(
        doc,
        ["Command", "Purpose"],
        [
            ["docker compose up -d --build", "Build and start all services."],
            ["docker compose ps", "Confirm frontend, backend, postgres, redis, and adminer are running."],
            ["docker compose logs --tail 80 backend", "Check backend errors or login/API activity."],
            ["docker compose down", "Stop the stack while keeping database data."],
            ["docker compose down -v", "Reset all database data. Use only when a fresh database is required."],
        ],
        widths=[2.7, 3.8],
    )

    doc.add_heading("7. Demo Scenario for Submission", level=1)
    doc.add_paragraph("Run the attack simulation after the stack is running and login works:")
    add_callout(
        doc,
        "Demo command",
        "powershell.exe -ExecutionPolicy Bypass -File .\\tools\\simulate-attacks.ps1 -BaseUrl http://127.0.0.1:3001",
    )
    add_steps(
        doc,
        [
            "Run the demo script and copy the Run ID shown in PowerShell.",
            "Open Dashboard and show updated threat metrics and charts.",
            "Open Logs and search the Run ID to show ingested events.",
            "Open Alerts and show detections for brute force, SQL injection, suspicious login, and credential dumping.",
            "Acknowledge an alert to show investigation started.",
            "Resolve an alert to show the issue is handled.",
            "Open AI Analysis and explain offline analysis mode if no real API key is configured.",
            "Open Adminer database GUI and show logs, alerts, users, rules, incidents, and audit records.",
        ],
    )

    doc.add_heading("8. Database GUI and Important Tables", level=1)
    add_table(
        doc,
        ["Adminer field", "Value"],
        [
            ["URL", "http://127.0.0.1:8080"],
            ["System", "PostgreSQL"],
            ["Server", "postgres"],
            ["Username", "siem_user"],
            ["Password", "password"],
            ["Database", "siem_db"],
        ],
        widths=[1.8, 4.7],
    )
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ["users", "Application users, role, active status, and password rotation timestamp."],
            ["logs", "Security events ingested into the SIEM."],
            ["alerts", "Threat detections generated by enabled rules."],
            ["incidents", "Analyst investigation and response records."],
            ["rules", "Detection rules and conditions."],
            ["log_sources", "Allowed log sources and API keys."],
            ["audit_log", "Login, profile, password, and other important user actions."],
            ["ai_analyses", "AI/offline analysis request history and output."],
        ],
        widths=[1.6, 4.9],
    )
    add_callout(doc, "Database safety", "Do not click Drop, Truncate, or Delete during submission unless you intentionally want to remove data.")

    doc.add_heading("9. Account and Password Management", level=1)
    doc.add_paragraph(
        "The Settings page supports access management for the logged-in user. The user can update full name, update email, change password, and log out. Passwords are stored as bcrypt hashes, not plain text."
    )
    add_bullets(
        doc,
        [
            "Default account: admin@siem.local / changeme123.",
            "Password change requires current password, new password, and confirmation.",
            "The system tracks password_changed_at in the users table.",
            "The Settings page shows the 15-day password rotation status.",
            "Logout returns the user to the login screen.",
        ],
    )

    doc.add_heading("10. AI Analysis Mode", level=1)
    doc.add_paragraph(
        "The project can use a real Anthropic API key, but it is not required for the lab. If ANTHROPIC_API_KEY is empty, the application uses deterministic offline SOC analysis mode."
    )
    add_table(
        doc,
        ["Mode", "When used", "How to explain"],
        [
            ["Offline mode", "No ANTHROPIC_API_KEY in .env", "Safe lab mode that still extracts attack type, indicators, MITRE hints, and response steps."],
            ["Online mode", "A real Anthropic API key is configured", "Calls Claude through the backend for live AI analysis."],
        ],
        widths=[1.5, 2.2, 2.8],
    )

    doc.add_heading("11. Common Issues and Fixes", level=1)
    add_table(
        doc,
        ["Issue", "Fix"],
        [
            ["Docker CLI permission denied", "Open Docker Desktop and run commands from the real Windows account."],
            ["localhost does not open in Codex browser", "Use http://127.0.0.1:5173."],
            ["Login failed on 127.0.0.1", "Backend CORS must allow both localhost and 127.0.0.1."],
            ["PowerShell blocks npm scripts", "Use npm.cmd run build."],
            ["Need database GUI", "Open Adminer at http://127.0.0.1:8080."],
            ["Need fresh database", "Use docker compose down -v, then rebuild and seed. This deletes existing data."],
        ],
        widths=[2.4, 4.1],
    )

    doc.add_heading("12. Viva Explanation Script", level=1)
    add_callout(
        doc,
        "Short answer",
        "This is an AI-assisted SIEM platform. It collects logs, stores them in PostgreSQL, applies detection rules, creates alerts, supports incidents, and provides SOC dashboard views for investigation.",
    )
    add_bullets(
        doc,
        [
            "Acknowledge means the analyst has started investigating an alert.",
            "Resolve means the alert has been handled and closed.",
            "Threats Blocked is a dashboard metric showing detected/handled malicious activity.",
            "Offline AI mode is used because a real paid API key is not required for the lab.",
            "Adminer proves the backend data is actually stored in PostgreSQL.",
        ],
    )

    doc.add_heading("13. Final Submission Checklist", level=1)
    add_bullets(
        doc,
        [
            "Docker Desktop is running.",
            "docker compose ps shows all containers running.",
            "Frontend opens at http://127.0.0.1:5173.",
            "Login works with admin@siem.local.",
            "Demo attack simulation has been run.",
            "Logs and alerts are visible in the UI.",
            "Adminer opens and shows PostgreSQL tables.",
            "Settings shows access, login/logout, password change, and about project.",
        ],
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
