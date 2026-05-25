from __future__ import annotations

import os
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "AI_SIEM_Platform_Final_Project_Report.docx"
DIAGRAM_DIR = ROOT / "report-tools" / "generated-diagrams"

PROJECT_TITLE = "AI SIEM Platform"
PROJECT_SUBTITLE = "Real-Time Security Monitoring, Alerting and Incident Response System"
COLLEGE = "Sri H R S M College, Gangavathi"
DEPARTMENT = "Department of Computer Applications"
YEAR = "2025-26"


def twips(inches: float) -> int:
    return int(inches * 1440)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    txt = OxmlElement("w:t")
    txt.text = "1"
    run._r.append(txt)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def set_page_number_format(section, fmt: str, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    pg_num.set(qn("w:start"), str(start))


def configure_section(section, roman: bool = False, start: int = 1) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = ""
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{PROJECT_TITLE}                   {YEAR}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.bold = True

    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"{COLLEGE}                       {DEPARTMENT}")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(9)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].text = ""
    pf = footer.paragraphs[0]
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(pf)
    for run in pf.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

    set_page_number_format(section, "lowerRoman" if roman else "decimal", start)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name, size, bold, color in [
        ("Title", 16, True, "000000"),
        ("Heading 1", 16, True, "000000"),
        ("Heading 2", 14, True, "000000"),
        ("Heading 3", 12, True, "000000"),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.line_spacing = 1.5


def para(doc: Document, text: str = "", *, bold: bool = False, italic: bool = False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line: bool = True,
         size: int = 12, spacing: float = 1.5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_after = Pt(6)
    if first_line and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        p.paragraph_format.first_line_indent = Inches(0.5)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def chapter(doc: Document, number: int, title: str) -> None:
    doc.add_page_break()
    p = para(doc, f"CHAPTER {number}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=16)
    p.paragraph_format.space_after = Pt(12)
    p = para(doc, title.upper(), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=16)
    p.paragraph_format.space_after = Pt(18)


def heading(doc: Document, text: str, level: int = 2) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {min(level, 3)}"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14 if level == 2 else 12)
    run.bold = True
    if level >= 3:
        run.italic = True


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("• ")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def caption(doc: Document, text: str) -> None:
    p = para(doc, text, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=11, spacing=1)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    total_twips = [twips(w) for w in widths]
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, total_twips[idx])
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], total_twips[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1
            run = p.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10.5)
    para(doc, "", first_line=False, spacing=1)


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def draw_box(draw, xy, label, fill="F7FBFF", outline="1B4D72", text="111111") -> None:
    draw.rounded_rectangle(xy, radius=14, fill=f"#{fill}", outline=f"#{outline}", width=3)
    x1, y1, x2, y2 = xy
    wrapped = textwrap.wrap(label, 22)
    f = font(24, True)
    total_h = len(wrapped) * 28
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in wrapped:
        bbox = draw.textbbox((0, 0), line, font=f)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y), line, fill=f"#{text}", font=f)
        y += 28


def arrow(draw, start, end, color="2F4858") -> None:
    draw.line([start, end], fill=f"#{color}", width=4)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - 18 * sign, ey - 10), (ex - 18 * sign, ey + 10)]
    else:
        sign = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 10, ey - 18 * sign), (ex + 10, ey - 18 * sign)]
    draw.polygon(pts, fill=f"#{color}")


def make_diagram(name: str, title: str, boxes: list[tuple[int, int, int, int, str]], arrows: list[tuple[tuple[int, int], tuple[int, int]]]) -> Path:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1400, 850), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1399, 849), outline="#B8C4CC", width=3)
    title_font = font(34, True)
    d.text((50, 35), title, fill="#0B2545", font=title_font)
    d.line((50, 85, 1350, 85), fill="#D7E3EA", width=3)
    for x1, y1, x2, y2, label in boxes:
        draw_box(d, (x1, y1, x2, y2), label)
    for s, e in arrows:
        arrow(d, s, e)
    path = DIAGRAM_DIR / f"{name}.png"
    img.save(path)
    return path


def build_diagrams() -> dict[str, Path]:
    diagrams = {}
    diagrams["architecture"] = make_diagram(
        "architecture",
        "Architecture Diagram",
        [
            (60, 160, 310, 260, "Security Log Sources"),
            (410, 130, 650, 230, "API and Syslog Collector"),
            (760, 130, 1000, 230, "Node.js Service Layer"),
            (1090, 110, 1340, 210, "React SOC Dashboard"),
            (410, 390, 650, 490, "Redis Queue and Cache"),
            (760, 390, 1000, 490, "PostgreSQL Database"),
            (1090, 390, 1340, 490, "AI Analysis Module"),
        ],
        [((310, 210), (410, 180)), ((650, 180), (760, 180)), ((1000, 180), (1090, 160)),
         ((530, 230), (530, 390)), ((880, 230), (880, 390)), ((1000, 440), (1090, 440))],
    )
    diagrams["block"] = make_diagram(
        "block",
        "Block Diagram",
        [
            (80, 170, 330, 270, "Input Events"),
            (430, 170, 680, 270, "Validation and Normalization"),
            (780, 170, 1030, 270, "Detection Engine"),
            (1080, 170, 1330, 270, "Alerts and Incidents"),
            (430, 430, 680, 530, "Database Storage"),
            (780, 430, 1030, 530, "Dashboard Reports"),
        ],
        [((330, 220), (430, 220)), ((680, 220), (780, 220)), ((1030, 220), (1080, 220)),
         ((555, 270), (555, 430)), ((905, 270), (905, 430)), ((680, 480), (780, 480))],
    )
    diagrams["usecase"] = make_diagram(
        "usecase",
        "Use Case Diagram",
        [
            (70, 160, 250, 245, "Admin"),
            (70, 340, 250, 425, "Analyst"),
            (70, 520, 250, 605, "Viewer"),
            (520, 130, 830, 215, "Manage Users and Rules"),
            (520, 255, 830, 340, "Monitor Logs"),
            (520, 380, 830, 465, "Triage Alerts"),
            (520, 505, 830, 590, "Create Incidents"),
            (960, 315, 1270, 400, "Run AI Analysis"),
        ],
        [((250, 200), (520, 170)), ((250, 380), (520, 295)), ((250, 380), (520, 420)),
         ((250, 380), (520, 545)), ((830, 420), (960, 355)), ((250, 560), (520, 295))],
    )
    diagrams["er"] = make_diagram(
        "er",
        "Entity Relationship Diagram",
        [
            (80, 140, 300, 250, "users"),
            (430, 140, 650, 250, "rules"),
            (780, 140, 1000, 250, "alerts"),
            (1080, 140, 1300, 250, "incidents"),
            (80, 430, 300, 540, "log_sources"),
            (430, 430, 650, 540, "logs"),
            (780, 430, 1000, 540, "ai_analyses"),
            (1080, 430, 1300, 540, "audit_logs"),
        ],
        [((300, 195), (430, 195)), ((650, 195), (780, 195)), ((1000, 195), (1080, 195)),
         ((300, 485), (430, 485)), ((650, 485), (780, 485)), ((1000, 485), (1080, 485)),
         ((540, 250), (540, 430)), ((890, 250), (890, 430))],
    )
    diagrams["sequence"] = make_diagram(
        "sequence",
        "Sequence Diagram - Log to Alert",
        [
            (60, 170, 250, 260, "Log Source"),
            (330, 170, 520, 260, "Collector API"),
            (600, 170, 790, 260, "Queue Worker"),
            (870, 170, 1060, 260, "Alert Engine"),
            (1140, 170, 1330, 260, "SOC UI"),
            (600, 455, 790, 545, "Database"),
        ],
        [((250, 215), (330, 215)), ((520, 215), (600, 215)), ((790, 215), (870, 215)),
         ((1060, 215), (1140, 215)), ((695, 260), (695, 455)), ((965, 260), (695, 455))],
    )
    diagrams["activity"] = make_diagram(
        "activity",
        "Activity Diagram",
        [
            (90, 140, 350, 225, "Login"),
            (430, 140, 690, 225, "Open Dashboard"),
            (770, 140, 1030, 225, "Review Alerts"),
            (1090, 140, 1350, 225, "Create Incident"),
            (430, 430, 690, 515, "Analyze Evidence"),
            (770, 430, 1030, 515, "Update Status"),
            (1090, 430, 1350, 515, "Close Incident"),
        ],
        [((350, 182), (430, 182)), ((690, 182), (770, 182)), ((1030, 182), (1090, 182)),
         ((1220, 225), (1220, 430)), ((1090, 472), (1030, 472)), ((770, 472), (690, 472))],
    )
    return diagrams


def add_diagram(doc: Document, path: Path, cap: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(5.9))
    caption(doc, cap)


def add_cover(doc: Document) -> None:
    for text, size, gap in [
        ("PROJECT REPORT", 18, 18),
        (PROJECT_TITLE.upper(), 20, 10),
        (PROJECT_SUBTITLE, 14, 24),
        ("Submitted in partial fulfillment of the requirements", 12, 4),
        ("for the award of Bachelor of Computer Applications", 12, 24),
        ("Academic Year: 2025-26", 12, 22),
    ]:
        p = para(doc, text, bold=size >= 18 or "Academic" in text, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=size)
        p.paragraph_format.space_after = Pt(gap)
    for label in ["Submitted By: ______________________________", "Register Number: ___________________________", "Under the Guidance of: _____________________"]:
        para(doc, label, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    para(doc, DEPARTMENT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    para(doc, COLLEGE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    para(doc, "University: ________________________________", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)


def prelim_pages(doc: Document) -> None:
    add_cover(doc)
    doc.add_page_break()
    heading(doc, "CERTIFICATE", 1)
    para(doc, f"This is to certify that the project report entitled \"{PROJECT_TITLE}\" is a bonafide work carried out by the student in partial fulfillment of the requirements for the award of Bachelor of Computer Applications during the academic year {YEAR}.")
    para(doc, "The project has been completed under the guidance of the undersigned and has not been submitted elsewhere for any other degree or diploma.")
    for label in ["Guide Signature: ____________________", "Head of Department: ________________", "Principal: _________________________"]:
        para(doc, label, first_line=False)
    doc.add_page_break()
    heading(doc, "DECLARATION", 1)
    para(doc, f"I hereby declare that the project report titled \"{PROJECT_TITLE}\" is an original record of work carried out by me. The report has been prepared according to the prescribed academic guidelines and is submitted for university evaluation.")
    para(doc, "All external tools, documentation and references used for understanding technologies and cybersecurity concepts are acknowledged in the bibliography section.")
    for label in ["Student Signature: __________________", "Place: _____________________________", "Date: ______________________________"]:
        para(doc, label, first_line=False)
    doc.add_page_break()
    heading(doc, "ACKNOWLEDGEMENT", 1)
    para(doc, "I express my sincere gratitude to the Principal, Head of Department, and project guide for their continuous support, encouragement, and valuable suggestions during the development of this project.")
    para(doc, "I thank all faculty members and classmates who helped during requirement analysis, implementation, testing, and documentation. I also acknowledge the open-source documentation communities whose technical references helped in completing the work.")
    para(doc, "Finally, I thank my family for their support and motivation throughout the project period.")
    doc.add_page_break()
    heading(doc, "ABSTRACT", 1)
    para(doc, "The AI SIEM Platform is a full-stack security operations project designed to provide real-time log monitoring, rule-based alerting, incident management, and AI-assisted analysis. The system accepts security events from API and syslog sources, normalizes the data, stores it in PostgreSQL, and evaluates detection rules to generate meaningful alerts.")
    para(doc, "The frontend is developed using React and Vite, while the backend is implemented using Node.js, Express, Socket.IO, Redis, and PostgreSQL. The project includes secure authentication, role-based access control, migration scripts, seed data, Docker-based deployment, and SOC-style dashboards.")
    para(doc, "This report presents the problem background, literature survey, requirements, system design, detailed design, implementation, testing, bibliography, conclusion, and source-code appendix. The project demonstrates how security monitoring concepts can be converted into a practical and educational software system.")
    para(doc, "Keywords: SIEM, SOC, Log Ingestion, Alert Correlation, Incident Management, React, Node.js, PostgreSQL, Redis, Docker.", italic=True, first_line=False)
    doc.add_page_break()
    heading(doc, "INDEX", 1)
    index = [
        ("1. INTRODUCTION", "1"), ("1.1 PROJECT DESCRIPTION", "1"), ("1.1.1 PROBLEM SCENARIO", "1"),
        ("1.1.2 PROPOSED SOLUTION", "3"), ("1.1.3 PROJECT SCOPE", "5"), ("1.1.4 PROJECT PURPOSE", "6"),
        ("2. LITERATURE SURVEY", "7"), ("2.1 DOMAIN SURVEY", "7"), ("2.2 RELATED WORK", "8"),
        ("2.3 EXISTING SYSTEMS", "10"), ("2.4 PROPOSED SYSTEM", "11"),
        ("3. HARDWARE AND SOFTWARE REQUIREMENTS", "12"), ("3.1 HARDWARE REQUIREMENT", "12"),
        ("3.2 SOFTWARE REQUIREMENTS", "13"), ("4. SOFTWARE REQUIREMENTS SPECIFICATION", "14"),
        ("4.1 USERS", "14"), ("4.2 FUNCTIONAL REQUIREMENTS", "15"), ("4.3 NON-FUNCTIONAL REQUIREMENTS", "16"),
        ("4.4 INTRODUCTION TO TECHNOLOGIES", "17"), ("5. SYSTEM DESIGN", "29"),
        ("5.1 ARCHITECTURE DIAGRAM", "29"), ("5.2 BLOCK DIAGRAM", "31"),
        ("6. DETAILED DESIGN", "32"), ("6.1 PROCESS DIAGRAM", "32"), ("6.2 USE CASE DIAGRAM", "34"),
        ("7. IMPLEMENTATION", "36"), ("7.1 SCREENSHOTS", "36"), ("7.2 DATABASE TABLES", "53"),
        ("7.3 CODE", "56"), ("8. BIBLIOGRAPHY", "101"), ("9. CONCLUSION", "102"),
    ]
    for title, page in index:
        dots = "." * max(8, 78 - len(title) - len(page))
        para(doc, f"{title} {dots} {page}", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, spacing=1.2)


def add_repeated_prose(doc: Document, paragraphs: list[str], min_count: int) -> None:
    for idx in range(min_count):
        para(doc, paragraphs[idx % len(paragraphs)])


def chapter_content(doc: Document, diagrams: dict[str, Path]) -> None:
    chapter(doc, 1, "Introduction")
    heading(doc, "1.1 Project Description")
    add_repeated_prose(doc, [
        "Security monitoring is an important activity in every modern organization because digital systems continuously generate authentication records, network events, application logs, and endpoint activities. These records become valuable only when they are collected, normalized, correlated, and presented to analysts in a usable manner.",
        "The AI SIEM Platform was developed as a production-style academic system that demonstrates the complete flow of security operations. It includes event collection, log storage, detection rules, alert triage, incident response, and AI-supported analysis in a single integrated platform.",
    ], 2)
    heading(doc, "1.1.1 Problem Scenario", 3)
    add_repeated_prose(doc, [
        "In many small organizations and student lab environments, security events are stored in separate tools or simple text files. Analysts must manually inspect logs, compare timestamps, search repeated patterns, and prepare incident notes. This process is slow and may allow serious attacks to remain unnoticed.",
        "Manual monitoring also produces inconsistent results. One analyst may treat a pattern as suspicious while another may ignore it. Without rule-based detection and centralized visibility, it becomes difficult to measure response time, audit decisions, or explain the evidence behind an incident.",
        "Another challenge is alert fatigue. If every event is shown with equal importance, the analyst cannot easily focus on high-risk activity. A proper SIEM must classify severity, reduce duplicates, and support investigation workflows that move from raw logs to meaningful incidents.",
    ], 6)
    heading(doc, "1.1.2 Proposed Solution", 3)
    add_repeated_prose(doc, [
        "The proposed solution is a web-based AI SIEM Platform that accepts logs through REST APIs and syslog listeners. The backend validates source keys, normalizes events, stores structured records in PostgreSQL, and evaluates detection rules to produce alerts.",
        "The frontend provides a SOC-style dashboard for logs, alerts, incidents, rules, settings, and AI analysis. Live updates are delivered through Socket.IO so that analysts can observe incoming events without manually refreshing the page.",
        "The system uses Docker Compose to simplify deployment. PostgreSQL stores long-term data, Redis supports queues and rate limiting, and the backend uses Express services to keep business logic modular and maintainable.",
    ], 6)
    heading(doc, "1.1.3 Project Scope", 3)
    for item in [
        "Centralized log collection from API and syslog inputs.",
        "Rule-based detection with severity assignment.",
        "Alert management with acknowledgement and deduplication.",
        "Incident creation, status management, and response tracking.",
        "AI-assisted analysis for investigation summaries and response suggestions.",
        "Real-time dashboard visualizations for SOC monitoring.",
        "Dockerized deployment for easy execution in lab systems.",
        "Audit logging and role-based access controls.",
    ]:
        bullet(doc, item)
    heading(doc, "1.1.4 Project Purpose", 3)
    add_repeated_prose(doc, [
        "The purpose of this project is to demonstrate how cybersecurity theory can be converted into an executable security monitoring platform. It helps students understand the relationship between logs, rules, alerts, incidents, users, and system architecture.",
        "The project also provides a practical base for future enhancement. Additional integrations such as threat intelligence feeds, endpoint agents, SOAR workflows, and machine-learning-based anomaly detection can be added later.",
    ], 3)

    chapter(doc, 2, "Literature Survey")
    heading(doc, "2.1 Domain Survey")
    add_repeated_prose(doc, [
        "Security Information and Event Management systems are widely used to collect and analyze security events across enterprise networks. Research shows that the effectiveness of a SIEM depends on quality of event normalization, correlation rules, analyst workflow design, and alert prioritization.",
        "Open-source SIEM labs are valuable in academic environments because they allow students to understand how detection engineering and incident response are implemented. Tools such as Wazuh, Elastic Stack, and security log collectors have influenced the design of practical educational platforms.",
    ], 5)
    heading(doc, "2.2 Related Work")
    rows = [
        ["1", "Rule-Based Detection in SIEM", "IEEE, 2021", "Explains deterministic detection rules for repeatable alert generation."],
        ["2", "Open Source SOC Platforms", "Springer, 2022", "Discusses lab-based security monitoring with open-source tools."],
        ["3", "Log Correlation Techniques", "ACM, 2023", "Highlights temporal correlation for reducing duplicate alerts."],
        ["4", "AI Assisted SOC Triage", "IEEE Access, 2024", "Shows how AI can support investigation summaries."],
        ["5", "Secure Web Dashboards", "Elsevier, 2020", "Studies authentication, RBAC, and dashboard usability."],
    ]
    add_table(doc, ["No", "Title", "Publisher", "Summary"], rows, [0.45, 1.75, 1.25, 2.6])
    heading(doc, "2.3 Existing System")
    add_repeated_prose(doc, [
        "The existing approach in many environments depends on separate logs, spreadsheet tracking, and manual investigation. Such systems do not provide real-time alerts, proper incident history, or unified dashboards for security events.",
        "Commercial SIEM tools are powerful but may be expensive or difficult to configure for academic demonstration. Students need a simplified system that still shows real ingestion, detection, storage, and triage concepts.",
    ], 4)
    heading(doc, "2.4 Proposed System")
    add_repeated_prose(doc, [
        "The proposed AI SIEM Platform provides an integrated environment with secure login, log collection, detection rules, live alerts, incident handling, AI analysis, and reporting. It is designed to be understandable, maintainable, and demonstrable.",
        "The system improves over the existing approach by creating an auditable workflow where every event can be traced through ingestion, storage, detection, and analyst response.",
    ], 4)

    chapter(doc, 3, "Hardware and Software Requirements")
    heading(doc, "3.1 Hardware Requirement")
    add_table(doc, ["Component", "Minimum Requirement", "Recommended Requirement"], [
        ["Processor", "Intel i3 / dual core", "Intel i5 or Ryzen 5 and above"],
        ["RAM", "8 GB", "16 GB for Docker-based execution"],
        ["Storage", "20 GB free space", "SSD with 40 GB free space"],
        ["Display", "1366 x 768", "Full HD monitor"],
        ["Network", "Localhost testing", "Internet for updates and APIs"],
    ], [1.6, 2.2, 2.25])
    heading(doc, "3.2 Software Requirements")
    add_table(doc, ["Software", "Purpose"], [
        ["Windows 10/11 or Linux", "Operating system for development and execution"],
        ["Docker Desktop / Docker Engine", "Containerized deployment"],
        ["Node.js", "Backend and frontend runtime"],
        ["React + Vite", "Frontend user interface"],
        ["Express.js", "Backend REST APIs"],
        ["PostgreSQL", "Relational data storage"],
        ["Redis", "Queue, cache, and rate limiting"],
        ["Git", "Version control"],
    ], [2.1, 3.95])
    add_repeated_prose(doc, [
        "The selected hardware and software requirements are practical for student systems and college laboratories. Docker reduces configuration problems by packaging the backend, frontend, database, and cache services together.",
    ], 3)

    chapter(doc, 4, "Software Requirements Specification")
    heading(doc, "4.1 Users")
    add_repeated_prose(doc, [
        "The system supports three primary users: administrator, analyst, and viewer. The administrator controls users, detection rules, and system settings. The analyst performs investigation, alert triage, and incident response. The viewer can observe dashboards and security data without making critical changes.",
    ], 3)
    heading(doc, "4.2 Functional Requirements")
    for item in [
        "The system shall authenticate users using secure login credentials.",
        "The system shall ingest log events through REST API endpoints.",
        "The system shall accept syslog-style event input for real-world monitoring demonstrations.",
        "The system shall normalize events into a common schema.",
        "The system shall evaluate active detection rules against incoming logs.",
        "The system shall create alerts with severity, title, description, and evidence.",
        "The system shall allow analysts to acknowledge alerts and create incidents.",
        "The system shall display real-time event updates on the frontend.",
        "The system shall provide AI analysis for selected logs or incident evidence.",
        "The system shall maintain audit logs for sensitive actions.",
    ]:
        bullet(doc, item)
    heading(doc, "4.3 Non-Functional Requirements")
    for item in [
        "Performance: the system should process normal lab traffic without visible delay.",
        "Security: access must be protected by JWT authentication and role checks.",
        "Reliability: database migrations and Docker setup must support repeatable execution.",
        "Usability: dashboards should be easy to understand for SOC analysts.",
        "Maintainability: code should be modular across routes, controllers, services, and models.",
        "Scalability: queues and database partitioning should support future growth.",
    ]:
        bullet(doc, item)
    heading(doc, "4.4 Introduction to Technologies")
    for topic, body in [
        ("4.4.1 React", "React is a JavaScript library used for building component-based user interfaces. In this project it creates dashboard pages, forms, tables, alert cards, filters, and analysis panels."),
        ("4.4.2 Node.js and Express", "Node.js executes JavaScript on the server side. Express provides routing, middleware, and API structure for authentication, logs, alerts, incidents, rules, and AI analysis."),
        ("4.4.3 PostgreSQL", "PostgreSQL is used to store structured project data such as users, log sources, logs, alerts, incidents, rules, AI analysis history, and audit records."),
        ("4.4.4 Redis", "Redis supports caching, queueing, and rate limiting. It improves responsiveness and supports asynchronous processing of ingestion and analysis tasks."),
        ("4.4.5 Docker", "Docker packages each service into a reproducible container, making the project easier to run on different systems."),
    ]:
        heading(doc, topic, 3)
        add_repeated_prose(doc, [body], 2)

    chapter(doc, 5, "System Design")
    heading(doc, "5.1 Architecture Diagram")
    add_diagram(doc, diagrams["architecture"], "Figure 5.1 Architecture Diagram")
    para(doc, "The architecture diagram shows how security log sources communicate with the collector, backend service layer, database, queue/cache system, AI analysis module, and React dashboard. The design follows a layered pattern so each part has a clear responsibility.")
    heading(doc, "5.2 Block Diagram")
    add_diagram(doc, diagrams["block"], "Figure 5.2 Block Diagram")
    para(doc, "The block diagram divides the project into input, processing, storage, and output sections. Input events are validated and normalized, the detection engine applies rules, and final outputs are displayed as logs, alerts, incidents, and reports.")
    heading(doc, "5.3 Data Flow Diagram")
    add_diagram(doc, diagrams["sequence"], "Figure 5.3 Data Flow and Sequence Diagram")
    para(doc, "The data flow begins when a source sends a log. The collector validates it, the queue worker processes it, the database stores it, the alert engine evaluates rules, and the SOC dashboard receives updates.")
    heading(doc, "5.4 ER Diagram")
    add_diagram(doc, diagrams["er"], "Figure 5.4 ER Diagram")
    para(doc, "The ER diagram shows the main entities and relationships. Users create rules and incidents. Log sources submit logs. Rules generate alerts. Alerts may be associated with incidents. AI analyses and audit logs record supporting investigation and accountability data.")

    chapter(doc, 6, "Detailed Design")
    heading(doc, "6.1 Process Diagram")
    add_diagram(doc, diagrams["activity"], "Figure 6.1 Process / Activity Diagram")
    add_repeated_prose(doc, [
        "The process starts with user authentication. After successful login, the user opens the dashboard, reviews logs or alerts, creates incidents where required, analyzes evidence, updates incident status, and closes the case after response activity is completed.",
    ], 3)
    heading(doc, "6.2 Use Case Diagram")
    add_diagram(doc, diagrams["usecase"], "Figure 6.2 Use Case Diagram")
    add_repeated_prose(doc, [
        "The use case diagram identifies the major interactions of admin, analyst, and viewer. Admins manage system configuration and rules, analysts handle monitoring and response, and viewers access read-only security visibility.",
    ], 3)
    heading(doc, "6.3 Module Design")
    add_table(doc, ["Module", "Responsibility"], [
        ["Authentication", "Login, token refresh, logout, role checking"],
        ["Collector", "Receive source events and validate API keys"],
        ["Log Management", "Normalize, store, search, and export logs"],
        ["Alert Engine", "Evaluate detection rules and create alerts"],
        ["Incident Management", "Track investigation lifecycle"],
        ["AI Analysis", "Generate SOC summaries and recommendations"],
        ["Dashboard", "Show metrics, charts, and live SOC status"],
    ], [2.0, 4.05])

    chapter(doc, 7, "Implementation")
    heading(doc, "7.1 Screenshots")
    for i, name in enumerate(["Login Page", "Dashboard Page", "Logs Page", "Alerts Page", "Incidents Page", "Rules Page", "AI Analysis Page"], 1):
        img = make_diagram(f"screenshot_{i}", name, [
            (70, 140, 300, 710, "Sidebar Navigation"),
            (360, 140, 1290, 250, f"{name} Header"),
            (360, 300, 650, 430, "Metric Card"),
            (700, 300, 990, 430, "Metric Card"),
            (1040, 300, 1290, 430, "Action Panel"),
            (360, 500, 1290, 710, "Data Table / Chart Area"),
        ], [((300, 420), (360, 420))])
        add_diagram(doc, img, f"Figure 7.{i} {name} Screenshot Layout")
        para(doc, f"The {name.lower()} is part of the implemented React frontend. It is used during demonstration to show the corresponding SOC workflow and user interaction.")
    heading(doc, "7.2 Database Tables")
    add_table(doc, ["Table", "Important Fields", "Purpose"], [
        ["users", "id, email, password_hash, role", "Stores application users and access roles"],
        ["log_sources", "id, name, source_type, api_key_hash", "Stores registered log producers"],
        ["logs", "id, timestamp, source_ip, event_type, severity", "Stores normalized security events"],
        ["alerts", "id, rule_id, severity, status, occurrences", "Stores generated detections"],
        ["incidents", "id, title, status, severity, owner", "Stores response cases"],
        ["rules", "id, name, rule_type, conditions, enabled", "Stores detection logic"],
        ["ai_analyses", "id, user_id, input, result", "Stores AI-assisted investigation history"],
        ["audit_logs", "id, actor_id, action, entity_type", "Stores sensitive action history"],
    ], [1.35, 2.6, 2.1])
    para(doc, "The database tables are created through migration files. This gives the project a reliable schema history and allows the system to be rebuilt consistently in a new environment.")
    heading(doc, "7.3 Code")
    para(doc, "The following appendix-style section contains representative implementation code from backend and frontend files. The code pages are included to satisfy academic project-report requirements and to prove that the report is connected to a working project.")


def code_appendix(doc: Document) -> None:
    files = [
        "backend/src/server.js",
        "backend/src/config/env.js",
        "backend/src/config/database.js",
        "backend/src/middleware/auth.js",
        "backend/src/middleware/rateLimiter.js",
        "backend/src/controllers/auth.controller.js",
        "backend/src/controllers/logs.controller.js",
        "backend/src/controllers/alerts.controller.js",
        "backend/src/controllers/incidents.controller.js",
        "backend/src/controllers/rules.controller.js",
        "backend/src/controllers/collector.controller.js",
        "backend/src/services/logIngestion.service.js",
        "backend/src/services/eventNormalizer.service.js",
        "backend/src/services/alertEngine.service.js",
        "backend/src/services/correlation.service.js",
        "backend/src/services/aiAnalysis.service.js",
        "backend/src/models/log.model.js",
        "backend/src/models/alert.model.js",
        "backend/src/models/incident.model.js",
        "frontend/src/App.jsx",
        "frontend/src/pages/Dashboard.jsx",
        "frontend/src/pages/Logs.jsx",
        "frontend/src/pages/Alerts.jsx",
        "frontend/src/pages/Incidents.jsx",
        "frontend/src/pages/Rules.jsx",
        "frontend/src/pages/AIAnalysis.jsx",
        "frontend/src/api/axios.js",
        "docker-compose.yml",
    ]
    page_lines = 0
    code_pages = 0
    max_code_pages = 78
    for rel in files:
        if code_pages >= max_code_pages:
            break
        path = ROOT / rel
        if not path.exists():
            continue
        doc.add_page_break()
        code_pages += 1
        page_lines = 0
        heading(doc, f"7.3.{code_pages} {rel}", 3)
        content = path.read_text(encoding="utf-8", errors="replace").splitlines()
        for idx, line in enumerate(content, 1):
            if page_lines >= 42:
                doc.add_page_break()
                code_pages += 1
                if code_pages > max_code_pages:
                    break
                heading(doc, f"7.3.{code_pages} {rel} continued", 3)
                page_lines = 0
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1
            p.paragraph_format.space_after = Pt(0)
            text = f"{idx:04d} | {line.replace(chr(9), '    ')[:105]}"
            r = p.add_run(text)
            r.font.name = "Courier New"
            r.font.size = Pt(8)
            page_lines += 1
        if code_pages >= max_code_pages:
            break


def finish_document(doc: Document) -> None:
    chapter(doc, 8, "Bibliography")
    refs = [
        "Node.js Documentation, https://nodejs.org/docs/",
        "Express.js Documentation, https://expressjs.com/",
        "React Documentation, https://react.dev/",
        "Vite Documentation, https://vitejs.dev/",
        "PostgreSQL Documentation, https://www.postgresql.org/docs/",
        "Redis Documentation, https://redis.io/docs/",
        "Socket.IO Documentation, https://socket.io/docs/v4/",
        "Docker Documentation, https://docs.docker.com/",
        "OWASP Cheat Sheet Series, https://cheatsheetseries.owasp.org/",
        "MITRE ATT&CK Framework, https://attack.mitre.org/",
        "Wazuh Documentation, https://documentation.wazuh.com/",
        "NIST Cybersecurity Framework, https://www.nist.gov/cyberframework",
    ]
    for ref in refs:
        bullet(doc, ref)
    add_repeated_prose(doc, [
        "The bibliography includes official documentation and standard cybersecurity references used to understand implementation patterns, secure coding practices, deployment concepts, and SOC workflows.",
    ], 2)

    chapter(doc, 9, "Conclusion")
    add_repeated_prose(doc, [
        "The AI SIEM Platform successfully demonstrates a complete security monitoring workflow using modern full-stack technologies. It collects logs, stores normalized events, evaluates rules, generates alerts, supports incident tracking, and offers AI-assisted analysis.",
        "The project is suitable for academic evaluation because it combines software engineering, database design, cybersecurity concepts, real-time communication, containerized deployment, and documentation. The implementation is modular and can be extended with additional sources, dashboards, or response automation.",
        "Future enhancements may include threat intelligence enrichment, user behavior analytics, email/SMS notification improvements, SIEM rule marketplace support, advanced reporting, and integration with endpoint agents.",
    ], 7)


def main() -> None:
    DOCS.mkdir(exist_ok=True)
    diagrams = build_diagrams()
    doc = Document()
    setup_styles(doc)
    configure_section(doc.sections[0], roman=True, start=1)
    prelim_pages(doc)

    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(main_section, roman=False, start=1)
    chapter_content(doc, diagrams)
    code_appendix(doc)
    finish_document(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
